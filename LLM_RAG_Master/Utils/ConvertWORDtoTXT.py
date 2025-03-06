import re
import os

from docx import Document
from LLM_RAG_Master.Utils.my_text_splitter import MyTextSplitter

def ConvertWordToTxt(input_file, myTextSplitter: MyTextSplitter) -> list:
    title_0 = os.path.basename(input_file).split('.')[0]
    output_text = []
    doc = Document(input_file)

    title_1 = None
    title_2 = None
    title_3 = None

    for paragraph in doc.paragraphs:
        text_temp = re.sub('\s+', '', paragraph.text).strip()
        if text_temp != '':
            xml = paragraph._p.xml

            # outlineLvl 大纲等级
            if xml.find('<w:outlineLvl') >= 0:
                outlineLvl_start_index = xml.find('<w:outlineLvl')
                outlineLvl_end_index = xml.find('>', outlineLvl_start_index)
                outlineLvl_value = xml[outlineLvl_start_index:outlineLvl_end_index + 1]
                outlineLvl_value = re.search("\d+", outlineLvl_value).group()

                if outlineLvl_value == '1':
                    title_1 = text_temp
                    title_2 = None
                    title_3 = None
                elif outlineLvl_value == '2':
                    title_2 = text_temp
                    title_3 = None
                elif outlineLvl_value == '3':
                    title_3 = text_temp

            text_splits = myTextSplitter.split_text(text_temp)
            for text_split in text_splits:
                text_line =[f"一级标题: <{title_0}>, 二级标题: <{title_1}>, 三级标题: <{title_2}>, 四级标题: <{title_3}>", text_split]
                output_text.append(text_line)
    return output_text

# The commented-out version of the function:
# def ConvertWordToTxt(input_file, myTextSplitter: MyTextSplitter) -> list:
#     doc = Document(input_file)
#     output_text = []
#     for paragraph in doc.paragraphs:
#         text_temp = re.sub(r'\s+', '', paragraph.text).strip()
#         if text_temp != '':
#             text_splits = myTextSplitter.split_text(text_temp)
#             for text_split in text_splits:
#                 output_text.append(text_split)
#     return output_text

if __name__ == "__main__":
    myTextSplitter = MyTextSplitter()

    # 指定目标路径
    directory = "../knowledge/doc/"

    # 遍历目录中的所有文件
    for root, dirs, files in os.walk(directory):
        for file in files:
            # 使用正则过滤以".docx"为结尾的文件
            match = re.search(r'\.docx$', file)

            if match:
                word_path = os.path.join(root, file)
                word_basename = os.path.basename(word_path)

                # 获取文件名
                match = re.findall(r'[^.]+\.docx', word_basename)[0]
                txt_path = '../knowledge/txt/' + match + '.txt'
                output_text = ConvertWordToTxt(word_path, myTextSplitter)

                # Uncomment the following to write the output to a file
                # with open(txt_path, 'w', encoding='utf-8') as f:
                #     f.write('\n'.join(output_text))

                print('转换后的文本保存到: ', txt_path)
            else:
                print('未找到以.docx结尾的文件，请确认文件夹下是否存在对应的文件。')
